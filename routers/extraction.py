"""
routers/extraction.py
---------------------
APIRouter for file parsing, AI extraction graph execution, saving report observations, rematching, and confirming child matches.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
from datetime import datetime
from typing import Optional

from fastapi import APIRouter, File, HTTPException, UploadFile
from bson import ObjectId

from common import (
    ConfirmMatchRequest,
    ExtractRequest,
    RematchRequest,
    SaveRequest,
    jsonable,
)
from db import (
    get_children_collection,
    get_center_children_for_matching_by_name,
    get_reports_collection,
)
from pipeline import (
    _confidence,
    _fuzzy_best_match,
    _matched_entry,
    extraction_graph,
    save_graph,
)

router = APIRouter(prefix="/api", tags=["Extraction & Reports"])


def extract_text_from_image_with_gemini(data: bytes, mime_type: str) -> str:
    from google import genai
    api_key = os.getenv("GOOGLE_API_KEY")
    client = genai.Client(api_key=api_key)
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                genai.types.Part.from_bytes(
                    data=data,
                    mime_type=mime_type,
                ),
                "Extract all readable text from this image (such as handwritten or typed progress reports). Preserve paragraph breaks. Do not add any conversational text, explanations, or formatting comments; return only the exact extracted text."
            ]
        )
        return response.text or ""
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Image text extraction via Gemini failed: {exc}")


def transcribe_audio_with_gemini(data: bytes, mime_type: str) -> str:
    from google import genai
    api_key = os.getenv("GOOGLE_API_KEY")
    client = genai.Client(api_key=api_key)
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                genai.types.Part.from_bytes(
                    data=data,
                    mime_type=mime_type,
                ),
                "Transcribe this audio recording into text exactly as spoken. Detect the language automatically. If the audio is in Hindi, transcribe it in Devanagari script; if it is in Marathi, transcribe it in Devanagari script; if it is in English, transcribe it in English. Return only the transcription itself, with no conversational preamble or markdown formatting."
            ]
        )
        return response.text or ""
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Audio transcription via Gemini failed: {exc}")


def extract_text_from_docx_fallback(data: bytes) -> str:
    import zipfile
    import io
    from xml.etree import ElementTree
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            xml_content = z.read("word/document.xml")
        root = ElementTree.fromstring(xml_content)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraphs = []
        for p in root.findall(".//w:p", ns):
            texts = [t.text for t in p.findall(".//w:t", ns) if t.text]
            if texts:
                paragraphs.append("".join(texts))
        return "\n".join(paragraphs)
    except Exception as exc:
        raise ValueError(f"Fallback docx XML extraction failed: {exc}")


def _report_fingerprint(raw_report: str, title: str, date_str: str, center: str) -> str:
    normalized = re.sub(r"\s+", " ", raw_report.strip().lower())
    if not normalized:
        normalized = f"{title.strip().lower()}|{date_str}|{center.strip().lower()}"
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def _fmt_submission(doc: dict) -> str:
    when = doc.get("last_saved_at") or doc.get("first_saved_at")
    when_str = when.strftime("%d %b %Y, %H:%M UTC") if isinstance(when, datetime) else "earlier"
    coaches = ", ".join(doc.get("coaches") or []) or "unknown coaches"
    return (
        f'"{doc.get("report_title", "Untitled Report")}" '
        f'(report date {doc.get("report_date") or "unknown"}) was already saved '
        f"on {when_str} by {coaches}"
    )


@router.post("/extract-file")
async def extract_file(file: UploadFile = File(...)) -> dict:
    """
    Extract plain text from an uploaded report file (.txt / .md / .docx / .pdf / images / audio)
    so the frontend can pre-fill the raw-report textarea.
    """
    filename = (file.filename or "").lower()
    data = await file.read()
    if len(data) > 100 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large (max 100 MB).")

    text = ""
    try:
        if filename.endswith((".txt", ".md", ".text")):
            text = data.decode("utf-8", errors="replace")
        elif filename.endswith(".docx"):
            try:
                import io
                from docx import Document
                document = Document(io.BytesIO(data))
                parts = [p.text for p in document.paragraphs]
                for table in document.tables:
                    for row in table.rows:
                        parts.append(" | ".join(cell.text for cell in row.cells))
                text = "\n".join(parts)
            except Exception as docx_exc:
                try:
                    text = extract_text_from_docx_fallback(data)
                except Exception:
                    raise docx_exc
        elif filename.endswith(".pdf"):
            try:
                import io
                from pypdf import PdfReader
            except ImportError:
                raise HTTPException(
                    status_code=501,
                    detail="pypdf is not installed on the server (pip install pypdf).",
                )
            reader = PdfReader(io.BytesIO(data))
            text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
        elif filename.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp")):
            mime_type = "image/png" if filename.endswith(".png") else "image/jpeg" if filename.endswith((".jpg", ".jpeg")) else f"image/{filename.split('.')[-1]}"
            text = extract_text_from_image_with_gemini(data, mime_type)
        elif filename.endswith((".mp3", ".wav", ".m4a", ".ogg", ".flac", ".webm", ".aac")):
            ext = filename.split(".")[-1]
            mime_type = "audio/mpeg" if ext == "mp3" else "audio/wav" if ext == "wav" else f"audio/{ext}"
            text = transcribe_audio_with_gemini(data, mime_type)
        else:
            raise HTTPException(
                status_code=415,
                detail="Unsupported file type. Upload docx, pdf, txt, png, jpg, jpeg, mp3, wav, m4a, ogg, etc.",
            )
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Could not read file: {exc}")

    text = text.strip()
    if not text:
        raise HTTPException(
            status_code=422,
            detail="No text could be extracted from this file.",
        )
    return {"filename": file.filename, "text": text, "chars": len(text)}


@router.post("/extract")
def extract(req: ExtractRequest) -> dict:
    title = req.report_title or "Untitled Report"

    initial_state = {
        "raw_report": req.raw_report.strip(),
        "report_title": title,
        "report_date": req.report_date,
        "coaches": req.coaches,
        "selected_center_id": req.center_id,
        "selected_center_name": req.center_name,
        "identified_names": None,
        "_extracted_children": None,
        "matched_children": None,
        "save_results": None,
        "error": None,
    }
    result = extraction_graph.invoke(initial_state)

    extracted_center = result.get("selected_center_name") or ""
    extracted_title = result.get("report_title") or "Untitled Report"
    extracted_date = result.get("report_date") or ""

    if not req.force and extracted_center:
        report_hash = _report_fingerprint(req.raw_report, extracted_title, extracted_date, extracted_center)
        reports = get_reports_collection()
        dup = reports.find_one({"report_hash": report_hash, "center_name": extracted_center})
        if dup:
            raise HTTPException(
                status_code=409,
                detail=f"Duplicate report: {_fmt_submission(dup)}. "
                       "Use 'Process anyway' if this is intentional.",
            )

    return {
        "identified_names": result.get("identified_names") or [],
        "matched_children": jsonable(result.get("matched_children") or []),
        "extracted_center_name": result.get("selected_center_name"),
        "extracted_center_id": result.get("selected_center_id"),
        "extracted_report_title": result.get("report_title"),
        "extracted_report_date": result.get("report_date"),
        "extracted_coaches": result.get("coaches"),
        "error": result.get("error"),
    }


@router.post("/save")
def save(req: SaveRequest) -> dict:
    title = req.report_title or "Untitled Report"
    center_name = req.center_name or ""
    report_hash = _report_fingerprint(req.raw_report, title, req.report_date, center_name)
    reports = get_reports_collection()

    if not req.force:
        dup = reports.find_one({"report_hash": report_hash, "center_name": center_name})
        if dup:
            raise HTTPException(
                status_code=409,
                detail=f"Duplicate report: {_fmt_submission(dup)}. "
                       "Use 'Save anyway' if this is intentional.",
            )
        soft = reports.find_one(
            {
                "report_title": title,
                "report_date": req.report_date,
                "center_name": center_name,
            }
        )
        if soft:
            raise HTTPException(
                status_code=409,
                detail=f"Possible duplicate: {_fmt_submission(soft)}. "
                       "Use 'Save anyway' if this is a different report.",
            )

    state = {
        "report_title": title,
        "report_date": req.report_date,
        "coaches": req.coaches,
        "selected_center_id": req.center_id,
        "selected_center_name": req.center_name,
        "matched_children": req.matched_children,
        "report_hash": report_hash,
        "force_save": req.force,
    }
    result = save_graph.invoke(state)
    save_results = result.get("save_results") or []

    if req.matched_children:
        now = datetime.utcnow()
        try:
            reports.update_one(
                {"report_hash": report_hash, "center_name": center_name},
                {
                    "$set": {
                        "report_title": title,
                        "report_date": req.report_date,
                        "coaches": req.coaches,
                        "center_id": req.center_id,
                        "raw_report": req.raw_report,
                        "children_saved": [
                            r.get("name") for r in save_results if r.get("success")
                        ],
                        "last_saved_at": now,
                    },
                    "$setOnInsert": {"first_saved_at": now},
                    "$inc": {"submission_count": 1},
                },
                upsert=True,
            )
        except Exception:
            pass

    return {"save_results": jsonable(save_results)}


@router.post("/rematch")
def rematch(req: RematchRequest) -> dict:
    collection = get_children_collection()
    name = req.name.strip()
    if not name:
        return {"entry": None, "score": 0}

    pattern = re.compile(re.escape(name), re.IGNORECASE)
    query: dict = {"child_name": pattern}
    if req.center_name:
        query["balgruha_name"] = req.center_name
    doc = collection.find_one(query)
    if doc:
        return {"entry": jsonable(_matched_entry(name, doc, "exact", 100)), "score": 100}

    roster = (
        get_center_children_for_matching_by_name(req.center_name)
        if req.center_name else []
    )
    best_doc, score = _fuzzy_best_match(name, roster)
    conf = _confidence(score)
    if best_doc and conf in ("high", "medium"):
        return {"entry": jsonable(_matched_entry(name, best_doc, conf, score)), "score": score}
    return {"entry": None, "score": int(score)}


@router.post("/confirm-match")
def confirm_match(req: ConfirmMatchRequest) -> dict:
    collection = get_children_collection()
    try:
        doc = collection.find_one({"_id": ObjectId(req.db_id)})
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid db_id")
    if not doc:
        raise HTTPException(status_code=404, detail="Child not found")
    return {"entry": jsonable(_matched_entry(req.name, doc, "fuzzy_suggested", req.score))}
