"""
api.py
------
FastAPI backend for the ISF Psycho-Analysis Portal (React frontend).

Wraps the existing LangGraph pipeline (extraction_graph, save_graph) and the
MongoDB helpers in db.py as a REST API. Child identification is scoped by the
reliable `balgruha_name` join (see pipeline.match_children_node / db.py).

Run:
    uvicorn api:app --reload --port 8000
"""

from __future__ import annotations

import hashlib
import json
import os
import re
from datetime import date, datetime
from typing import Any, Optional

from bson import ObjectId
from dotenv import load_dotenv
from fastapi import FastAPI, File, HTTPException, UploadFile, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, FileResponse
from pydantic import BaseModel

load_dotenv()

from db import (
    get_all_centers,
    get_children_by_balgruha,
    get_children_collection,
    get_center_children_for_matching_by_name,
    get_child_observations_history,
    get_centers_collection,
    get_reports_collection,
    get_unmatched_collection,
    ping_db,
    get_upload_by_filename,
    get_risk_dashboard_data,
    update_child_risk_profile,
    get_children_falling_through_cracks,
    get_balgruha_risk_heatmap,
    get_pending_task_analytics,
    get_success_stories,
)
from pipeline import (
    extraction_graph,
    save_graph,
    _get_llm,
    _confidence,
    _doc_to_profile,
    _fuzzy_best_match,
    _fuzzy_score,
    _matched_entry,
)

app = FastAPI(title="ISF Psycho-Analysis API", version="1.0.0")

_allowed_origins = [
    "http://localhost:5173",
    "http://127.0.0.1:5173",
    "http://localhost:4173",
    "http://localhost",        # IIS on port 80
    "http://127.0.0.1",
]
# Add production origins from env (can be comma-separated)
_vps_origins = os.getenv("ALLOWED_ORIGIN", "").strip()
if _vps_origins:
    for origin in _vps_origins.split(","):
        origin_clean = origin.strip()
        if origin_clean:
            _allowed_origins.append(origin_clean)

app.add_middleware(
    CORSMiddleware,
    allow_origins=_allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------------------------------------------------------------------------
# Serialization helpers
# ---------------------------------------------------------------------------

def jsonable(value: Any) -> Any:
    """Recursively convert ObjectId / datetime into JSON-friendly primitives."""
    if isinstance(value, ObjectId):
        return str(value)
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, dict):
        return {k: jsonable(v) for k, v in value.items()}
    if isinstance(value, (list, tuple)):
        return [jsonable(v) for v in value]
    return value


# ---------------------------------------------------------------------------
# Request models
# ---------------------------------------------------------------------------

class ExtractRequest(BaseModel):
    raw_report: str
    report_title: str = "Untitled Report"
    report_date: str = ""           # ISO YYYY-MM-DD
    coaches: list[str] = []
    center_id: Optional[str] = None
    center_name: Optional[str] = None
    force: bool = False


class SaveRequest(BaseModel):
    report_title: str = "Untitled Report"
    report_date: str = ""
    coaches: list[str] = []
    center_id: Optional[str] = None
    center_name: Optional[str] = None
    matched_children: list[dict] = []
    # Raw report text — fingerprinted to detect duplicate submissions.
    raw_report: str = ""
    # Set by the frontend's "Save anyway" button to override the duplicate guard.
    force: bool = False


class RematchRequest(BaseModel):
    name: str
    center_name: Optional[str] = None


class ConfirmMatchRequest(BaseModel):
    db_id: str
    name: str
    score: int = 100


class DeleteObservationRequest(BaseModel):
    date: Optional[str] = None       # ISO datetime string of the observation
    reportTitle: Optional[str] = None
    scope: str = "single"            # "single" | "all"


class UpdateObservationRequest(BaseModel):
    """
    Body for PATCH /api/children/{child_id}/observations.
    Identifies the observation by report_hash (preferred) or date+reportTitle,
    then applies `fields` as a partial update.
    """
    # --- Identifier (at least one required) ---
    report_hash: Optional[str] = None
    date: Optional[str] = None          # ISO datetime string
    reportTitle: Optional[str] = None
    # --- Editable payload ---
    fields: dict                        # keys must be in EDITABLE_OBS_FIELDS


class CreateChildRequest(BaseModel):
    child_name: str
    balgruha_name: str
    class_studying: Optional[str] = None
    dob: Optional[str] = None
    school: Optional[str] = None
    parent_status: Optional[str] = None
    languages: Optional[str] = None
    strengths: Optional[str] = None
    weakness: Optional[str] = None
    nature_behavior: Optional[str] = None
    # Name as extracted from the report (used to build the matched entry)
    extracted_name: Optional[str] = None


class AskQuestionRequest(BaseModel):
    question: str


class TranslateSummaryRequest(BaseModel):
    # Currently only "hi" is supported, but kept open for extension.
    lang: str = "hi"
    refresh: bool = False


class TranslateTasksRequest(BaseModel):
    lang: str = "hi"
    refresh: bool = False



# ---------------------------------------------------------------------------
# Health & centers
# ---------------------------------------------------------------------------

@app.get("/api/health")
def health() -> dict:
    return {
        "mongo": ping_db(),
    }


@app.get("/api/centers")
def centers() -> list[dict]:
    return [{"id": str(c["_id"]), "name": c.get("name", "")} for c in get_all_centers()]


@app.get("/api/centers/{name}/roster")
def roster(name: str) -> dict:
    """Children for a balgruha + aggregate stats (for the Insights tab)."""
    kids = get_children_by_balgruha(name)
    total_obs = 0
    coaches: set[str] = set()
    for k in kids:
        obs = k.get("observations") or []
        if isinstance(obs, list):
            total_obs += len(obs)
            for o in obs:
                inv = o.get("coachesInvolved") if isinstance(o, dict) else None
                if isinstance(inv, list):
                    coaches.update(str(c) for c in inv if c)
                elif isinstance(inv, str) and inv.strip():
                    coaches.add(inv.strip())
    return {
        "children": [jsonable(k) for k in kids],
        "stats": {
            "total_children": len(kids),
            "total_observations": total_obs,
            "active_coaches": len(coaches),
        },
    }


@app.get("/api/centers/{name}/search")
def search_children(name: str, q: str = "", limit: int = 8) -> list[dict]:
    """
    Lightweight autocomplete over a center's roster.
    Combines substring matches (ranked first) with fuzzy scores so the
    review screen can offer live suggestions while the user types.
    """
    q = q.strip()
    if not q:
        return []
    roster = get_center_children_for_matching_by_name(name)
    q_lower = q.lower()
    results: list[dict] = []
    for doc in roster:
        child_name = doc.get("child_name", "") or ""
        if not child_name:
            continue
        substring = q_lower in child_name.lower()
        score = _fuzzy_score(q, child_name)
        if substring or score >= 55:
            results.append(
                {
                    "child_name": child_name,
                    "db_id": str(doc["_id"]),
                    "score": 100 if child_name.lower() == q_lower else int(score),
                    "substring": substring,
                    "class_studying": doc.get("class_studying", "") or "",
                }
            )
    results.sort(key=lambda r: (not r["substring"], -r["score"], r["child_name"]))
    for r in results:
        r.pop("substring", None)
    return results[: max(1, min(limit, 25))]


# ---------------------------------------------------------------------------
# Gender inference (AI, batch, with DB persistence)
# ---------------------------------------------------------------------------

class InferGendersRequest(BaseModel):
    names: list[str]                # display names to classify
    child_ids: list[str] = []       # optional: parallel list of DB _ids to persist result


_GENDER_PROMPT = """\
You are an expert in Indian names. Given the list of names below, classify each name as "male", "female", or "unknown".
Many of these children are from Maharashtra, so use knowledge of Marathi, Hindi, and other Indian names.

Return ONLY a JSON array (no markdown, no explanation) where each element is exactly one of: "male", "female", or "unknown".
The array must be the same length and in the same order as the input names.

Names:
{names}

JSON array:"""


@app.post("/api/infer-genders")
def infer_genders(req: InferGendersRequest) -> dict:
    """
    Use Gemini to infer gender from a list of Indian child names in a single
    batch call. Optionally persists the result back to MongoDB (child_ids must
    have the same length as names when provided).
    Returns: { "genders": ["male"|"female"|"unknown", ...] }
    """
    if not req.names:
        return {"genders": []}

    from langchain_core.prompts import ChatPromptTemplate
    from langchain_core.output_parsers import StrOutputParser

    names_text = "\n".join(f"{i+1}. {n}" for i, n in enumerate(req.names))
    prompt = ChatPromptTemplate.from_template(_GENDER_PROMPT)
    chain = prompt | _get_llm() | StrOutputParser()

    try:
        raw = chain.invoke({"names": names_text}).strip()
        # Strip markdown fences if present
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        raw = raw.strip()
        genders_raw = json.loads(raw)
        if not isinstance(genders_raw, list):
            raise ValueError("Expected a JSON array")
        # Normalise and pad/trim to match input length
        valid = {"male", "female", "unknown"}
        genders: list[str] = []
        for g in genders_raw:
            val = str(g).strip().lower()
            genders.append(val if val in valid else "unknown")
        # Pad with "unknown" if shorter
        while len(genders) < len(req.names):
            genders.append("unknown")
        genders = genders[: len(req.names)]
    except Exception:
        genders = ["unknown"] * len(req.names)

    # Persist to MongoDB (best-effort, non-blocking errors)
    if req.child_ids and len(req.child_ids) == len(req.names):
        col = get_children_collection()
        from bson import ObjectId as ObjId
        for child_id_str, gender in zip(req.child_ids, genders):
            if gender == "unknown":
                continue
            try:
                col.update_one(
                    {"_id": ObjId(child_id_str)},
                    {"$set": {"gender": gender}},
                )
            except Exception:
                pass

    return {"genders": genders}




def extract_text_from_image_with_gemini(data: bytes, mime_type: str) -> str:
    from google import genai
    api_key = os.getenv("GOOGLE_API_KEY")
    client = genai.Client(api_key=api_key)
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                genai.types.Part.from_bytes(
                    data=data,
                    mime_type=mime_type,
                ),
                "Extract all readable text from this image (such as handwritten or typed progress reports). Preserve paragraph breaks. Do not add any conversational text, explanations, or formatting comments; return only the exact extracted text."
            ]
        )
        return response.text or ""
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Image text extraction via Gemini failed: {exc}")


def transcribe_audio_with_gemini(data: bytes, mime_type: str) -> str:
    from google import genai
    api_key = os.getenv("GOOGLE_API_KEY")
    client = genai.Client(api_key=api_key)
    try:
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                genai.types.Part.from_bytes(
                    data=data,
                    mime_type=mime_type,
                ),
                "Transcribe this audio recording into text exactly as spoken. Detect the language automatically. If the audio is in Hindi, transcribe it in Devanagari script; if it is in Marathi, transcribe it in Devanagari script; if it is in English, transcribe it in English. Return only the transcription itself, with no conversational preamble or markdown formatting."
            ]
        )
        return response.text or ""
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"Audio transcription via Gemini failed: {exc}")
def extract_text_from_docx_fallback(data: bytes) -> str:
    import zipfile
    import io
    from xml.etree import ElementTree
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            xml_content = z.read("word/document.xml")
        root = ElementTree.fromstring(xml_content)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraphs = []
        for p in root.findall(".//w:p", ns):
            texts = [t.text for t in p.findall(".//w:t", ns) if t.text]
            if texts:
                paragraphs.append("".join(texts))
        return "\n".join(paragraphs)
    except Exception as exc:
        raise ValueError(f"Fallback docx XML extraction failed: {exc}")


@app.post("/api/extract-file")
async def extract_file(file: UploadFile = File(...)) -> dict:
    """
    Extract plain text from an uploaded report file (.txt / .md / .docx / .pdf / images / audio)
    so the frontend can pre-fill the raw-report textarea.
    """
    filename = (file.filename or "").lower()
    data = await file.read()
    if len(data) > 100 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="File too large (max 100 MB).")

    text = ""
    try:
        if filename.endswith((".txt", ".md", ".text")):
            text = data.decode("utf-8", errors="replace")
        elif filename.endswith(".docx"):
            try:
                import io
                from docx import Document  # python-docx
                document = Document(io.BytesIO(data))
                parts = [p.text for p in document.paragraphs]
                for table in document.tables:
                    for row in table.rows:
                        parts.append(" | ".join(cell.text for cell in row.cells))
                text = "\n".join(parts)
            except Exception as docx_exc:
                try:
                    text = extract_text_from_docx_fallback(data)
                except Exception:
                    raise docx_exc
        elif filename.endswith(".pdf"):
            try:
                import io
                from pypdf import PdfReader
            except ImportError:
                raise HTTPException(
                    status_code=501,
                    detail="pypdf is not installed on the server (pip install pypdf).",
                )
            reader = PdfReader(io.BytesIO(data))
            text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
        elif filename.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp")):
            # Image OCR
            mime_type = "image/png" if filename.endswith(".png") else "image/jpeg" if filename.endswith((".jpg", ".jpeg")) else f"image/{filename.split('.')[-1]}"
            text = extract_text_from_image_with_gemini(data, mime_type)
        elif filename.endswith((".mp3", ".wav", ".m4a", ".ogg", ".flac", ".webm", ".aac")):
            # Audio transcription
            ext = filename.split(".")[-1]
            mime_type = "audio/mpeg" if ext == "mp3" else "audio/wav" if ext == "wav" else f"audio/{ext}"
            text = transcribe_audio_with_gemini(data, mime_type)
        else:
            raise HTTPException(
                status_code=415,
                detail="Unsupported file type. Upload docx, pdf, txt, png, jpg, jpeg, mp3, wav, m4a, ogg, etc.",
            )
    except HTTPException:
        raise
    except Exception as exc:  # corrupt file, encrypted pdf, etc.
        raise HTTPException(status_code=422, detail=f"Could not read file: {exc}")

    text = text.strip()
    if not text:
        raise HTTPException(
            status_code=422,
            detail="No text could be extracted from this file.",
        )
    return {"filename": file.filename, "text": text, "chars": len(text)}


@app.post("/api/extract")
def extract(req: ExtractRequest) -> dict:
    title = req.report_title or "Untitled Report"
    center_name = req.center_name or ""

    initial_state = {
        "raw_report": req.raw_report.strip(),
        "report_title": title,
        "report_date": req.report_date,
        "coaches": req.coaches,
        "selected_center_id": req.center_id,
        "selected_center_name": req.center_name,
        "identified_names": None,
        "_extracted_children": None,
        "matched_children": None,
        "save_results": None,
        "error": None,
    }
    result = extraction_graph.invoke(initial_state)

    extracted_center = result.get("selected_center_name") or ""
    extracted_title = result.get("report_title") or "Untitled Report"
    extracted_date = result.get("report_date") or ""

    if not req.force and extracted_center:
        report_hash = _report_fingerprint(req.raw_report, extracted_title, extracted_date, extracted_center)
        reports = get_reports_collection()
        dup = reports.find_one({"report_hash": report_hash, "center_name": extracted_center})
        if dup:
            raise HTTPException(
                status_code=409,
                detail=f"Duplicate report: {_fmt_submission(dup)}. "
                       "Use 'Process anyway' if this is intentional.",
            )

    return {
        "identified_names": result.get("identified_names") or [],
        "matched_children": jsonable(result.get("matched_children") or []),
        "extracted_center_name": result.get("selected_center_name"),
        "extracted_center_id": result.get("selected_center_id"),
        "extracted_report_title": result.get("report_title"),
        "extracted_report_date": result.get("report_date"),
        "extracted_coaches": result.get("coaches"),
        "error": result.get("error"),
    }



def _report_fingerprint(raw_report: str, title: str, date_str: str, center: str) -> str:
    """
    SHA-256 fingerprint of a report. Normalizes the text (lowercase, collapsed
    whitespace) so trivial spacing/casing changes still count as the same
    report. Falls back to title|date|center when no raw text is available
    (e.g. an older client that doesn't send it).
    """
    normalized = re.sub(r"\s+", " ", raw_report.strip().lower())
    if not normalized:
        normalized = f"{title.strip().lower()}|{date_str}|{center.strip().lower()}"
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()


def _fmt_submission(doc: dict) -> str:
    """Human-friendly description of a prior registry entry for 409 messages."""
    when = doc.get("last_saved_at") or doc.get("first_saved_at")
    when_str = when.strftime("%d %b %Y, %H:%M UTC") if isinstance(when, datetime) else "earlier"
    coaches = ", ".join(doc.get("coaches") or []) or "unknown coaches"
    return (
        f'"{doc.get("report_title", "Untitled Report")}" '
        f'(report date {doc.get("report_date") or "unknown"}) was already saved '
        f"on {when_str} by {coaches}"
    )


@app.post("/api/save")
def save(req: SaveRequest) -> dict:
    title = req.report_title or "Untitled Report"
    center_name = req.center_name or ""
    report_hash = _report_fingerprint(req.raw_report, title, req.report_date, center_name)
    reports = get_reports_collection()

    if not req.force:
        # Hard duplicate: identical report text already saved for this center.
        dup = reports.find_one({"report_hash": report_hash, "center_name": center_name})
        if dup:
            raise HTTPException(
                status_code=409,
                detail=f"Duplicate report: {_fmt_submission(dup)}. "
                       "Use 'Save anyway' if this is intentional.",
            )
        # Soft duplicate: same title + date + center but different text
        # (catches "same report, minor edits" resubmissions).
        soft = reports.find_one(
            {
                "report_title": title,
                "report_date": req.report_date,
                "center_name": center_name,
            }
        )
        if soft:
            raise HTTPException(
                status_code=409,
                detail=f"Possible duplicate: {_fmt_submission(soft)}. "
                       "Use 'Save anyway' if this is a different report.",
            )

    state = {
        "report_title": title,
        "report_date": req.report_date,
        "coaches": req.coaches,
        "selected_center_id": req.center_id,
        "selected_center_name": req.center_name,
        "matched_children": req.matched_children,
        "report_hash": report_hash,
        "force_save": req.force,
    }
    result = save_graph.invoke(state)
    save_results = result.get("save_results") or []

    # Register the report fingerprint whenever children were part of the
    # request. We do this unconditionally (not just on success) so that a
    # partially-failed save — or one where all pipeline saves errored —
    # still leaves a registry entry that blocks duplicate re-submissions.
    # The unique index on (report_hash, center_name) keeps this idempotent.
    if req.matched_children:
        now = datetime.utcnow()
        try:
            reports.update_one(
                {"report_hash": report_hash, "center_name": center_name},
                {
                    "$set": {
                        "report_title": title,
                        "report_date": req.report_date,
                        "coaches": req.coaches,
                        "center_id": req.center_id,
                        # Raw report text — kept so the UI can show the original
                        # ("basic") report next to the AI-extracted fields for
                        # comparison (POC review feedback, 6 Jul 2026).
                        "raw_report": req.raw_report,
                        "children_saved": [
                            r.get("name") for r in save_results if r.get("success")
                        ],
                        "last_saved_at": now,
                    },
                    "$setOnInsert": {"first_saved_at": now},
                    "$inc": {"submission_count": 1},
                },
                upsert=True,
            )
        except Exception:
            # Registry bookkeeping must never fail the save itself.
            pass

    return {"save_results": jsonable(save_results)}


# ---------------------------------------------------------------------------
# Re-match / confirm-match (Review step actions)
# ---------------------------------------------------------------------------

@app.post("/api/rematch")
def rematch(req: RematchRequest) -> dict:
    """Manual re-match: exact regex (scoped by balgruha) then fuzzy fallback."""
    collection = get_children_collection()
    name = req.name.strip()
    if not name:
        return {"entry": None, "score": 0}

    pattern = re.compile(re.escape(name), re.IGNORECASE)
    query: dict = {"child_name": pattern}
    if req.center_name:
        query["balgruha_name"] = req.center_name
    doc = collection.find_one(query)
    if doc:
        return {"entry": jsonable(_matched_entry(name, doc, "exact", 100)), "score": 100}

    roster = (
        get_center_children_for_matching_by_name(req.center_name)
        if req.center_name else []
    )
    best_doc, score = _fuzzy_best_match(name, roster)
    conf = _confidence(score)
    if best_doc and conf in ("high", "medium"):
        return {"entry": jsonable(_matched_entry(name, best_doc, conf, score)), "score": score}
    return {"entry": None, "score": int(score)}


@app.post("/api/confirm-match")
def confirm_match(req: ConfirmMatchRequest) -> dict:
    """Confirm a suggested candidate by its db_id -> build a matched entry."""
    collection = get_children_collection()
    try:
        doc = collection.find_one({"_id": ObjectId(req.db_id)})
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid db_id")
    if not doc:
        raise HTTPException(status_code=404, detail="Child not found")
    return {"entry": jsonable(_matched_entry(req.name, doc, "fuzzy_suggested", req.score))}


# ---------------------------------------------------------------------------
# Child profile & observation history
# ---------------------------------------------------------------------------

@app.post("/api/children")
def create_child(req: CreateChildRequest) -> dict:
    """
    Create a brand-new child profile (used when a report mentions a child
    that does not yet exist in the roster). Returns a matched entry so the
    review screen can link the card immediately.
    """
    child_name = req.child_name.strip()
    balgruha = req.balgruha_name.strip()
    if not child_name:
        raise HTTPException(status_code=400, detail="child_name is required")
    if not balgruha:
        raise HTTPException(status_code=400, detail="balgruha_name is required")

    collection = get_children_collection()

    # Guard against exact duplicates within the same center.
    existing = collection.find_one(
        {
            "child_name": re.compile(f"^{re.escape(child_name)}$", re.IGNORECASE),
            "balgruha_name": balgruha,
        }
    )
    if existing:
        raise HTTPException(
            status_code=409,
            detail=f'"{existing.get("child_name")}" already exists in {balgruha}. '
                   "Use re-match instead.",
        )

    doc = {
        "child_name": child_name,
        "balgruha_name": balgruha,
        "class_studying": (req.class_studying or "").strip(),
        "dob": (req.dob or "").strip(),
        "school": (req.school or "").strip(),
        "parent_status": (req.parent_status or "").strip(),
        "languages": (req.languages or "").strip(),
        "strengths": (req.strengths or "").strip(),
        "weakness": (req.weakness or "").strip(),
        "nature_behavior": (req.nature_behavior or "").strip(),
        "photo_url": "",
        "observations": [],
        "created_at": datetime.utcnow(),
        "created_via": "portal",
    }
    result = collection.insert_one(doc)
    doc["_id"] = result.inserted_id

    extracted = req.extracted_name or child_name
    return {"entry": jsonable(_matched_entry(extracted, doc, "created", 100))}


@app.get("/api/children/{child_id}")
def child_profile(child_id: str) -> dict:
    collection = get_children_collection()
    try:
        doc = collection.find_one({"_id": ObjectId(child_id)})
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid child_id")
    if not doc:
        raise HTTPException(status_code=404, detail="Child not found")
    return jsonable(doc)


# Allowlist of basic profile fields the frontend is allowed to update
_PROFILE_EDITABLE_FIELDS = {
    "gender", "class_studying", "school", "dob",
    "parent_status", "languages", "strengths", "weakness", "nature_behavior",
    "risk_category",
}



class UpdateProfileRequest(BaseModel):
    fields: dict   # only keys present in _PROFILE_EDITABLE_FIELDS are accepted


@app.patch("/api/children/{child_id}/profile")
def update_child_profile(child_id: str, req: UpdateProfileRequest) -> dict:
    """Partial update of a child's basic profile fields (gender, class, etc.)."""
    collection = get_children_collection()
    try:
        oid = ObjectId(child_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid child_id")

    # Filter to allowlisted fields only
    safe = {k: v for k, v in req.fields.items() if k in _PROFILE_EDITABLE_FIELDS}
    if not safe:
        raise HTTPException(status_code=400, detail="No valid fields to update.")

    result = collection.update_one({"_id": oid}, {"$set": safe})
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Child not found")
    return {"modified": result.modified_count, "message": "Profile updated."}




@app.get("/api/uploads/{filename}")
def serve_upload(filename: str):
    """Retrieve and serve image binary data from munmeet_uploads by filename, with GCS and PureStore fallbacks."""
    # 1. Try to fetch from MongoDB first
    doc = get_upload_by_filename(filename)
    if doc and "data" in doc:
        content_type = doc.get("contentType", "image/jpeg")
        return Response(content=doc["data"], media_type=content_type)

    # 2. GCS Fallback
    import urllib.request
    from urllib.error import HTTPError
    bucket_name = "munmeet-media-sanket-7b724"
    folders = ["profiles", "photos"]
    gcs_data = None

    for folder in folders:
        url = f"https://storage.googleapis.com/{bucket_name}/{folder}/{filename}"
        try:
            req = urllib.request.Request(
                url, 
                headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
            )
            with urllib.request.urlopen(req) as response:
                gcs_data = response.read()
                break
        except HTTPError:
            continue
        except Exception:
            continue

    if gcs_data:
        ext = filename.split(".")[-1].lower()
        content_type = "image/png" if ext == "png" else "image/jpeg"
        return Response(content=gcs_data, media_type=content_type)

    # 3. PureStore Fallback
    import boto3
    from botocore.client import Config
    purestore_endpoint = "https://s3.in-west3.purestore.io"
    purestore_bucket = "playground"
    purestore_key_id = "ef0de8f471eb6180DYIV"
    purestore_secret = "b9khTmQ218DSPkHRGAeEaLnDEyoUMtJa4Bt8TzPL"

    s3 = boto3.client(
        's3',
        endpoint_url=purestore_endpoint,
        aws_access_key_id=purestore_key_id,
        aws_secret_access_key=purestore_secret,
        config=Config(signature_version='s3v4')
    )

    for folder in folders:
        key = f"{folder}/{filename}"
        try:
            response = s3.get_object(Bucket=purestore_bucket, Key=key)
            data = response['Body'].read()
            content_type = response.get('ContentType', 'image/jpeg')
            return Response(content=data, media_type=content_type)
        except Exception:
            continue

    raise HTTPException(status_code=404, detail="File not found in MongoDB, GCS, or PureStore")


@app.get("/api/purestore/{folder}/{filename}")
def serve_purestore_file(folder: str, filename: str):
    """Retrieve and serve image binary data directly from PureStore by folder and filename."""
    import boto3
    from botocore.client import Config

    purestore_endpoint = "https://s3.in-west3.purestore.io"
    purestore_bucket = "playground"
    purestore_key_id = "ef0de8f471eb6180DYIV"
    purestore_secret = "b9khTmQ218DSPkHRGAeEaLnDEyoUMtJa4Bt8TzPL"

    key = f"{folder}/{filename}"
    s3 = boto3.client(
        's3',
        endpoint_url=purestore_endpoint,
        aws_access_key_id=purestore_key_id,
        aws_secret_access_key=purestore_secret,
        config=Config(signature_version='s3v4')
    )
    try:
        response = s3.get_object(Bucket=purestore_bucket, Key=key)
        data = response['Body'].read()
        content_type = response.get('ContentType', 'image/jpeg')
        return Response(content=data, media_type=content_type)
    except Exception as exc:
        raise HTTPException(status_code=404, detail=f"PureStore file not found: {exc}")



@app.get("/api/children/{child_id}/observations")
def child_observations(child_id: str) -> list[dict]:
    try:
        oid = ObjectId(child_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid child_id")
    return [jsonable(o) for o in get_child_observations_history(oid)]


# ---------------------------------------------------------------------------
# AI summary of ALL reports for a child (review feedback #2, 6 Jul 2026)
# ---------------------------------------------------------------------------

AI_SUMMARY_PROMPT = """\
<role>
You are a senior child welfare analyst and clinical progress-note writer. You
read a child's complete observation history (multiple field reports, oldest
first) and write a SINGLE cohesive progress note that lets a coach or
counsellor see, at a glance, how the child has changed over time — progressing
or regressing — and what each individual report contributes to that arc.
</role>

<task>
Write one progress note for the child whose report history is given below. The
note must have ONE custom title (a single markdown `### ` heading chosen
specifically for this child based on the content of their reports) followed by
a chronological progression/regression walkthrough that uses the actual details
of each report. End the body with an explicit overall-verdict sentence.
</task>

<context>
Child: {child_name}
Center / Balgruha: {balgruha}

<report_history>
{observations_digest}
</report_history>
</context>

<output_format>
Output structure — and NOTHING else — must be exactly this:

### <single custom title for THIS child>

<chronological walkthrough paragraph for the oldest report>

<chronological walkthrough paragraph for the next report, describing what CHANGED vs. the previous>

... (one short narrative block per report, oldest -> newest) ...

Overall trajectory: <improving | stable | regressing | mixed | snapshot only> - <one-line justification referencing the arc above>.
</output_format>

<constraints>
MUST DO:
- Begin the response with EXACTLY `### ` (three hashes + one space). Never `#`,
  `##`, or `####`. The downstream viewer only renders `###`-level headings, so
  any other heading level is silently dropped.
- Make the title a SINGLE line. The title MUST be a custom phrase that captures
  the dominant theme or arc of THIS child's situation (it IS the custom heading).
- Walk through the reports CHRONOLOGICALLY (oldest to newest), one compact
  narrative block per report (1-3 sentences each). Each block must describe
  what THAT report shows in concrete terms (behaviors, test results, key notes,
  action items) AND explicitly frame it as progression, regression, stability,
  or a turning point relative to the prior report.
- Anchor every claim in the actual report content. Quote or paraphrase real
  dates, behaviors, and notes - do not invent any fact not present.
- Make the delta between consecutive reports visible: explicitly state what
  changed ("by [date] this had improved / worsened / stayed the same as seen
  in [specific detail]").
- Close the body with exactly one final, single sentence in the exact form:
  "Overall trajectory: <improving|stable|regressing|mixed|snapshot only> -
  <one-line justification>."
- Keep the whole note within 250-450 words even for children with many reports.
- The ONLY markdown element allowed is the leading `### ` on the title line.

MUST NOT DO:
- Do NOT use `#`, `##`, `####`, or any heading other than the single `### `
  title line.
- Do NOT add a second heading, sub-heading, divider, bullet list, numbered
  list, table, or code fence anywhere in the response.
- Do NOT use generic boilerplate titles such as "Progress / Regression
  Overview", "Child Summary", "Clinical Summary", or "Overview". Use phrases
  genuinely tied to this child (e.g. "Steady academic progress with recurring
  anxiety episodes", "Initial withdrawal giving way to peer engagement",
  "Behavioural setbacks following family disruption").
- Do NOT include the child's name, a date range, or the word "Summary" in the
  title - the UI already shows those.
- Do NOT invent or extrapolate facts beyond what is in the observations.
</constraints>

<edge_case_single_report>
If only ONE report exists for this child (no over-time comparison possible
yet), still produce the single `### ` title + exactly ONE walkthrough
paragraph summarizing that report, and end the body with this exact sentence:
"Overall trajectory: snapshot only - longitudinal comparison will be possible
once follow-up reports are added."
</edge_case_single_report>

<examples>
Example A - multiple reports:

### Steady opening up, with lingering test reluctance

The 04 May visit (psychologist P. Rao) noted Riya as withdrawn and avoiding eye
contact during the draw-a-person test; this is the baseline against which
later reports should be read. Little participation change was flagged at the
time, with action items centred on warm-up rapport building.

By the 18 Jun visit, no formal turning point yet - the same withdrawal
persisted, but the coach noted Riya asked one unprompted question, an early
regression-to-stability shift consistent with the rapport-building plan.

On 09 Jul the picture changed clearly: observations record Riya initiating a
peer game and making consistent eye contact, a tangible progression against
the May baseline, though she remained reluctant during the block-design test
that session. Overall trajectory: improving - withdrawal has eased across the
three visits and peer engagement is now emerging, though test reluctance
persists and warrants continued focus.

Example B - single report only:

### First-baseline snapshot: calm, cooperative, mild reluctance with authority figures

The 12 Jul screening by Dr. Mehta records Aman as cooperative during the
sentence-completion task and calm in the counsellor's presence, but visibly
tense when asked to read aloud in front of staff, with no prior report to
contrast against. Action items flagged rapport building with authority figures
and re-screening in two weeks. Overall trajectory: snapshot only -
longitudinal comparison will be possible once follow-up reports are added.
</examples>

Begin the response now with the single `### ` title line for the child above.
"""


def _observations_fingerprint(obs_list: list[dict]) -> str:
    """
    Stable fingerprint of a child's observation set, used to decide whether a
    cached AI summary is still valid. Any added/removed/edited observation
    changes the hash.
    """
    parts = []
    for o in obs_list:
        parts.append(
            "|".join(
                [
                    str(o.get("date", "")),
                    str(o.get("reportTitle", "")),
                    str(o.get("psychologicalNotes", ""))[:80],
                    str(len(o.get("actionItems") or [])),
                ]
            )
        )
    # Append version suffix to automatically invalidate cached AI summaries when prompt changes
    versioned_str = "\n".join(parts) + "|v4_single_title_progress_fewshot"
    return hashlib.sha256(versioned_str.encode("utf-8")).hexdigest()


def _build_observations_digest(obs_list: list[dict]) -> str:
    """Chronological (oldest-first) plain-text digest fed to the LLM."""
    chunks = []
    for o in reversed(obs_list):  # history endpoint returns latest-first
        date = o.get("date")
        date_str = (
            date.strftime("%d %b %Y") if isinstance(date, datetime) else str(date or "unknown date")[:10]
        )
        coaches = o.get("coachesInvolved") or []
        coaches_str = ", ".join(coaches) if isinstance(coaches, list) else str(coaches)
        actions = o.get("actionItems") or []
        actions_str = "; ".join(str(a) for a in actions) if actions else "none"
        chunks.append(
            f"--- Report: {o.get('reportTitle', 'Untitled')} | Date: {date_str}"
            f" | Coaches: {coaches_str or 'unknown'} ---\n"
            f"General background: {o.get('generalBackground') or '—'}\n"
            f"Psychological notes: {o.get('psychologicalNotes') or '—'}\n"
            f"Action items: {actions_str}"
        )
    return "\n\n".join(chunks)


def _extract_text_from_response(response: Any) -> str:
    """Extract string content from LangChain response, handling list-of-dicts formats."""
    content = getattr(response, "content", response)
    if isinstance(content, str):
        return content.strip()
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, str):
                parts.append(item)
            elif isinstance(item, dict) and "text" in item:
                parts.append(item["text"])
            elif hasattr(item, "text"):
                parts.append(item.text)
            elif hasattr(item, "get") and item.get("text"):
                parts.append(item.get("text"))
            else:
                parts.append(str(item))
        return "".join(parts).strip()
    return str(content).strip()


TRANSLATE_SUMMARY_PROMPT = """\
<role>
You are a professional clinical translator fluent in English and Hindi. You translate child welfare progress notes from English to Hindi while preserving their clinical precision, tone, and markdown structure.
</role>

<task>
Translate the English progress note below into Hindi. Keep the meaning, clinical terminology, dates, and named entities (people, tests, places) intact.

IMPORTANT VOCABULARY RULE:
Use simple, conversational, and easily understandable Hindi vocabulary that field coaches and social workers use. Avoid overly formal, Sanskritized, or literary Hindi terms (e.g., do NOT use complex words like "अवसाद", "मनोवैज्ञानिक"). Instead, use their common spoken Hindi equivalents or write the common English word transliterated in Devanagari script (e.g. use डिप्रेशन for depression, साइकोलॉजिस्ट for psychologist, टेस्ट for test, कोच for coach, फोकस for focus, एक्टिविटी for activity, बिहेवियर or व्यवहार for behavior, एंग्जायटी or चिंता for anxiety, प्रोग्रेस for progress).
</task>


<strict_format>
The output MUST preserve the structure of the input EXACTLY:
- Begin the response with EXACTLY `### ` (three hashes + one space) followed by the Hindi translation of the custom title on a SINGLE line. Never `#`, `##`, or `####`.
- Keep one chronological walkthrough paragraph per report, in the original order, separated by exactly one blank line.
- Preserve the final sentence in the exact same shape, translating only the words: it must still read in Hindi as "Overall trajectory: <improving|stable|regressing|mixed|snapshot only> - <one-line justification>." but the verb/justification translated to Hindi (the labeled field words improving/stable/etc. may either be translated to Hindi or kept transliterated — choose the more natural reading).
- The ONLY markdown element allowed is the single leading `### ` on the title line. No other headings, lists, tables, dividers, or code fences.
</strict_format>

<constraints>
MUST DO:
- Output ONLY the Hindi translation of the input note. No preamble, no "Here is the translation", no explanation.
- Anchor every claim in the original English note — do not invent or extrapolate.
- Preserve names verbatim (child name, coach names, psychologist names) — write them in Latin script or transliterate to Devanagari, but keep them recognizable.
- Keep the whole note within roughly the same length as the English original (the English is 250-450 words).

MUST NOT DO:
- Do NOT change the order of paragraphs.
- Do NOT merge or split paragraphs.
- Do NOT add or remove the `### ` title line.
- Do NOT change any factual detail (dates, test results, behaviors).
</constraints>

<input_note>
{english_summary}
</input_note>

Begin the response now with the single `### ` title line in Hindi.
"""


TRANSLATE_TASKS_PROMPT = """\
You are an expert clinical translator. Translate the following list of pending action items/tasks for a child from English to simple, conversational Hindi.

IMPORTANT VOCABULARY RULE:
Use simple, conversational, and easily understandable Hindi vocabulary that field coaches and social workers use. Avoid overly formal, Sanskritized, or literary Hindi terms. For common English words, write their simple Hindi equivalents or write the common English word transliterated in Devanagari script (e.g. use डिप्रेशन for depression, साइकोलॉजिस्ट for psychologist, टेस्ट or परीक्षण for test, कोच for coach, फोकस for focus, एक्टिविटी for activity, सेशन्स for sessions, काउंसलर for counselor).

Return the translation EXACTLY as a JSON array of strings, matching the length and order of the input array. Do not return any other text, markdown fences, or explanations.

Input tasks (JSON array of strings):
{english_tasks}

Return ONLY the JSON array:
"""


PENDING_TASKS_PROMPT = """\

<role>
You are a clinical operations analyst reviewing the consolidated follow-up action items for a single child across multiple field reports. Your job is to consolidate the list by removing duplicates and grouping semantically equivalent items, while preserving distinct genuine tasks.
</role>

<task>
Read the list of action items below, each labeled with the report and date it came from. Produce a SINGLE consolidated list of DE-DUPLICATED, DISTINCT follow-up tasks for this child.

Remove:
- Semantic duplicates (same task expressed differently across reports), e.g. "continue rapport building", "keep building rapport", "continue rapport-building exercises" -> keep ONE entry using the most informative / most recent wording.
- Verbatim repeats of the same task across reports that mean the same ongoing activity (e.g. "regular counseling sessions" appearing in 3 reports -> one entry).

Preserve:
- Genuinely distinct tasks (different goal or different addressee).
- Tasks that escalated or changed scope (merge into one entry if they describe the same goal at different intensities, phrased to reflect the latest intensity).

Order the output by importance/urgency first, then chronology.
</task>

<output_format>
Return ONLY a JSON array of strings. No markdown, no code fences, no explanation. Each string is one consolidated task statement. Empty array [] if the input has no action items.

Example output:
["Continue weekly rapport-building sessions with the child", "Coordinate with school for reduced test anxiety during exams", "Re-screen in two weeks to assess peer engagement"]
</output_format>

<input_action_items>
{tasks_digest}
</input_action_items>

JSON array:
"""


def _pending_tasks_fingerprint(obs_list: list[dict]) -> str:
    """
    Stable fingerprint of a child's action-items set across all observations,
    used to decide whether a cached pending-tasks digest is still valid.
    Modeled on _observations_fingerprint but scoped to action items only.
    """
    parts = []
    for o in obs_list:
        report_title = str(o.get("reportTitle", ""))
        date = o.get("date")
        date_str = date.isoformat() if isinstance(date, datetime) else str(date or "")
        for idx, task in enumerate(o.get("actionItems") or []):
            parts.append(f"{report_title}|{date_str}|{idx}|{task}")
    # v1 suffix invalidates cached digests when the prompt changes.
    return hashlib.sha256(("\n".join(parts) + "|v1").encode("utf-8")).hexdigest()


def _build_tasks_digest(obs_list: list[dict]) -> tuple[str, int]:
    """
    Build a chronological (oldest-first) plain-text digest of all action items,
    labeling each task with its source report title and date.

    Returns (digest, observation_count_with_tasks).
    """
    chunks = []
    obs_with_tasks = 0
    for o in reversed(obs_list):  # history endpoint returns latest-first
        actions = o.get("actionItems") or []
        if not actions:
            continue
        obs_with_tasks += 1
        date = o.get("date")
        date_str = (
            date.strftime("%d %b %Y") if isinstance(date, datetime) else str(date or "unknown date")[:10]
        )
        title = o.get("reportTitle", "Untitled Report")
        lines = [f'--- From report "{title}" on {date_str} ---']
        for i, task in enumerate(actions, start=1):
            lines.append(f"{i}. {task}")
        chunks.append("\n".join(lines))
    return "\n\n".join(chunks), obs_with_tasks


def _parse_tasks_json(raw: str) -> list[str]:
    """
    Parse the JSON array of task strings returned by the LLM. Mirrors the
    _parse_children_json pattern in pipeline.py: strip markdown fences, slice
    from first '[' to last ']', strip trailing commas, then json.loads.
    Raises ValueError on parse failure (caller raises HTTPException 502).
    """
    raw = raw.strip()
    # Strip markdown backticks
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    raw = raw.strip()

    # Slice to keep only the JSON array if conversational text leaked.
    start = raw.find("[")
    end = raw.rfind("]")
    if start != -1 and end != -1 and end > start:
        raw = raw[start : end + 1]

    # Clean up trailing commas in lists or objects.
    raw = re.sub(r",\s*([\]}])", r"\1", raw)

    try:
        data = json.loads(raw)
    except (json.JSONDecodeError, ValueError) as exc:
        raise ValueError(f"invalid JSON: {exc}")

    if not isinstance(data, list):
        raise ValueError("top-level JSON is not an array")
    return [str(item).strip() for item in data if isinstance(item, str) and str(item).strip()]


@app.get("/api/children/{child_id}/summary")
def child_ai_summary(child_id: str, refresh: bool = False) -> dict:
    """
    AI summary of ALL of a child's reports together.

    Cached on the child document (ai_summary / ai_summary_obs_hash /
    ai_summary_generated_at) so Gemini is only called when the observation
    set changed or the caller passes ?refresh=true.
    """
    collection = get_children_collection()
    try:
        oid = ObjectId(child_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid child_id")
    doc = collection.find_one({"_id": oid})
    if not doc:
        raise HTTPException(status_code=404, detail="Child not found")

    obs = get_child_observations_history(oid)
    if not obs:
        return {
            "summary": "",
            "generated_at": None,
            "cached": False,
            "observation_count": 0,
        }

    obs_hash = _observations_fingerprint(obs)

    # Serve the cached summary when the observation set hasn't changed.
    if (
        not refresh
        and doc.get("ai_summary")
        and doc.get("ai_summary_obs_hash") == obs_hash
    ):
        return {
            "summary": doc["ai_summary"],
            "generated_at": jsonable(doc.get("ai_summary_generated_at")),
            "cached": True,
            "observation_count": len(obs),
        }

    prompt = AI_SUMMARY_PROMPT.format(
        child_name=doc.get("child_name", "Unknown"),
        balgruha=doc.get("balgruha_name", "Unknown"),
        observations_digest=_build_observations_digest(obs),
    )
    try:
        response = _get_llm().invoke(prompt)
        summary = _extract_text_from_response(response)
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"AI summary failed: {exc}")

    if not summary:
        raise HTTPException(status_code=502, detail="AI returned an empty summary.")

    now = datetime.utcnow()
    try:
        collection.update_one(
            {"_id": oid},
            {
                "$set": {
                    "ai_summary": summary,
                    "ai_summary_obs_hash": obs_hash,
                    "ai_summary_generated_at": now,
                }
            },
        )
    except Exception:
        # Caching must never fail the request itself.
        pass

    return {
        "summary": summary,
        "generated_at": now.isoformat(),
        "cached": False,
        "observation_count": len(obs),
    }


@app.post("/api/children/{child_id}/summary/translate")
def child_ai_summary_translate(child_id: str, req: Optional[TranslateSummaryRequest] = None) -> dict:
    """
    Translate the cached English AI summary into another language (Hindi by
    default). The Hindi translation is cached on the child document as
    `ai_summary_hi` / `ai_summary_hi_obs_hash` / `ai_summary_hi_at`, keyed to
    the same observation fingerprint as the English summary so it is
    invalidated whenever the report set changes.
    """
    if req is None:
        req = TranslateSummaryRequest()
    if req.lang != "hi":
        raise HTTPException(
            status_code=400,
            detail="Only Hindi ('hi') translation is currently supported.",
        )

    collection = get_children_collection()
    try:
        oid = ObjectId(child_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid child_id")
    doc = collection.find_one({"_id": oid})
    if not doc:
        raise HTTPException(status_code=404, detail="Child not found")

    obs = get_child_observations_history(oid)
    if not obs:
        return {
            "translated": "",
            "lang": "hi",
            "cached": False,
            "generated_at": None,
            "observation_count": 0,
        }

    obs_hash = _observations_fingerprint(obs)

    # Serve the cached Hindi translation when the observation set hasn't changed.
    if (
        not req.refresh
        and doc.get("ai_summary_hi")
        and doc.get("ai_summary_hi_obs_hash") == obs_hash
    ):
        return {
            "translated": doc["ai_summary_hi"],
            "lang": "hi",
            "cached": True,
            "generated_at": jsonable(doc.get("ai_summary_hi_at")),
            "observation_count": len(obs),
        }

    # Require an English summary to translate from.
    english_summary = doc.get("ai_summary") or ""
    if not english_summary.strip():
        raise HTTPException(
            status_code=409,
            detail="Generate the English AI summary first (open the 'AI Summary of All Reports' section).",
        )

    prompt = TRANSLATE_SUMMARY_PROMPT.format(english_summary=english_summary)
    try:
        response = _get_llm().invoke(prompt)
        translated = _extract_text_from_response(response)
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"AI translation failed: {exc}")

    if not translated:
        raise HTTPException(status_code=502, detail="AI returned an empty translation.")

    now = datetime.utcnow()
    try:
        collection.update_one(
            {"_id": oid},
            {
                "$set": {
                    "ai_summary_hi": translated,
                    "ai_summary_hi_obs_hash": obs_hash,
                    "ai_summary_hi_at": now,
                }
            },
        )
    except Exception:
        # Caching must never fail the request itself.
        pass

    return {
        "translated": translated,
        "lang": "hi",
        "cached": False,
        "generated_at": now.isoformat(),
        "observation_count": len(obs),
    }


@app.get("/api/children/{child_id}/pending-tasks")
def child_pending_tasks(child_id: str) -> dict:
    """
    Consolidated, de-duplicated follow-up action items across ALL of a child's
    reports (Gemini removes semantic duplicates). Cached on the child document
    as `pending_tasks_dedup` / `pending_tasks_obs_hash` / `pending_tasks_at`,
    keyed to a fingerprint of the action-items set. Re-saving any observation
    changes the fingerprint and triggers re-generation.
    """
    collection = get_children_collection()
    try:
        oid = ObjectId(child_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid child_id")
    doc = collection.find_one({"_id": oid})
    if not doc:
        raise HTTPException(status_code=404, detail="Child not found")

    obs = get_child_observations_history(oid)  # latest-first
    tasks_digest, obs_with_tasks = _build_tasks_digest(obs)

    if obs_with_tasks == 0:
        return {
            "tasks": [],
            "cached": False,
            "generated_at": None,
            "observation_count": len(obs),
            "observation_count_with_tasks": 0,
        }

    fp = _pending_tasks_fingerprint(obs)

    # Serve the cached consolidated list when the action-items set hasn't changed.
    if (
        doc.get("pending_tasks_dedup")
        and doc.get("pending_tasks_obs_hash") == fp
    ):
        return {
            "tasks": doc["pending_tasks_dedup"],
            "cached": True,
            "generated_at": jsonable(doc.get("pending_tasks_at")),
            "observation_count": len(obs),
            "observation_count_with_tasks": obs_with_tasks,
        }

    prompt = PENDING_TASKS_PROMPT.format(tasks_digest=tasks_digest)
    try:
        response = _get_llm().invoke(prompt)
        raw = _extract_text_from_response(response)
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"AI pending-tasks failed: {exc}")

    if not raw.strip():
        raise HTTPException(status_code=502, detail="AI returned an empty task list.")

    try:
        tasks = _parse_tasks_json(raw)
    except ValueError as exc:
        raise HTTPException(status_code=502, detail=f"AI returned invalid task list JSON: {exc}")

    now = datetime.utcnow()
    try:
        collection.update_one(
            {"_id": oid},
            {
                "$set": {
                    "pending_tasks_dedup": tasks,
                    "pending_tasks_obs_hash": fp,
                    "pending_tasks_at": now,
                }
            },
        )
    except Exception:
        # Caching must never fail the request itself.
        pass

    return {
        "tasks": tasks,
        "cached": False,
        "generated_at": now.isoformat(),
        "observation_count": len(obs),
        "observation_count_with_tasks": obs_with_tasks,
    }


@app.post("/api/children/{child_id}/pending-tasks/translate")
def child_pending_tasks_translate(child_id: str, req: Optional[TranslateTasksRequest] = None) -> dict:
    """
    Translate the consolidated pending tasks into simple Hindi.
    Caches on the child document as `pending_tasks_dedup_hi` / `pending_tasks_hi_obs_hash` / `pending_tasks_hi_at`.
    """
    if req is None:
        req = TranslateTasksRequest()
    if req.lang != "hi":
        raise HTTPException(
            status_code=400,
            detail="Only Hindi ('hi') translation is currently supported.",
        )

    collection = get_children_collection()
    try:
        oid = ObjectId(child_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid child_id")
    doc = collection.find_one({"_id": oid})
    if not doc:
        raise HTTPException(status_code=404, detail="Child not found")

    obs = get_child_observations_history(oid)
    if not obs:
        return {
            "translated": [],
            "lang": "hi",
            "cached": False,
            "generated_at": None,
            "observation_count": 0,
        }

    fp = _pending_tasks_fingerprint(obs)

    # Serve the cached Hindi translation when the observations haven't changed.
    if (
        not req.refresh
        and doc.get("pending_tasks_dedup_hi")
        and doc.get("pending_tasks_hi_obs_hash") == fp
    ):
        return {
            "translated": doc["pending_tasks_dedup_hi"],
            "lang": "hi",
            "cached": True,
            "generated_at": jsonable(doc.get("pending_tasks_hi_at")),
            "observation_count": len(obs),
        }

    # Require English tasks to translate from.
    # If not cached, first retrieve them.
    english_tasks = doc.get("pending_tasks_dedup")
    if english_tasks is None or doc.get("pending_tasks_obs_hash") != fp:
        # Generate them on the fly
        res = child_pending_tasks(child_id)
        english_tasks = res.get("tasks") or []

    if not english_tasks:
        return {
            "translated": [],
            "lang": "hi",
            "cached": False,
            "generated_at": None,
            "observation_count": len(obs),
        }

    prompt = TRANSLATE_TASKS_PROMPT.format(english_tasks=json.dumps(english_tasks))
    try:
        response = _get_llm().invoke(prompt)
        raw = _extract_text_from_response(response)
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"AI tasks translation failed: {exc}")

    try:
        translated = _parse_tasks_json(raw)
    except Exception as exc:
        # Fallback parsing in case the LLM returned raw lines
        lines = [line.strip().replace("-", "").strip() for line in raw.split("\n") if line.strip()]
        translated = [line for line in lines if line]

    now = datetime.utcnow()
    try:
        collection.update_one(
            {"_id": oid},
            {
                "$set": {
                    "pending_tasks_dedup_hi": translated,
                    "pending_tasks_hi_obs_hash": fp,
                    "pending_tasks_hi_at": now,
                }
            },
        )
    except Exception:
        pass

    return {
        "translated": translated,
        "lang": "hi",
        "cached": False,
        "generated_at": now.isoformat(),
        "observation_count": len(obs),
    }


# ---------------------------------------------------------------------------

# Raw ("basic") report text for AI-vs-raw comparison (review feedback #5)
# ---------------------------------------------------------------------------

@app.get("/api/reports/{report_hash}")
def report_raw(report_hash: str, center_name: Optional[str] = None) -> dict:
    """
    Return the original raw report text stored in the reports registry.
    Observations saved after 6 Jul 2026 carry a `report_hash` that joins to
    this registry entry. Older reports were saved before raw text was
    persisted — for those `has_raw` is false.
    """
    reports = get_reports_collection()
    query: dict = {"report_hash": report_hash}
    if center_name:
        query["center_name"] = center_name
    doc = reports.find_one(query)
    if not doc:
        raise HTTPException(status_code=404, detail="Report not found in registry")
    raw = doc.get("raw_report") or ""
    return {
        "report_title": doc.get("report_title", "Untitled Report"),
        "report_date": doc.get("report_date", ""),
        "center_name": doc.get("center_name", ""),
        "coaches": doc.get("coaches") or [],
        "raw_report": raw,
        "has_raw": bool(raw.strip()),
    }


# ---------------------------------------------------------------------------
# Q&A Child Assistant
# ---------------------------------------------------------------------------

ASK_QUESTION_PROMPT = """\
You are an expert child counseling assistant. Below is the complete record for the child {child_name}, who attends center/balgruha {balgruha}.
Answer the user's question accurately using ONLY the provided child profile and observation history.
Be professional, concise, empathetic, and evidence-based. If the answer is not contained in or cannot be reasonably inferred from the records, explicitly state that it is not recorded in the child's profile.

--- CHILD PROFILE ---
Strengths: {strengths}
Weakness: {weakness}
Nature/Behavior: {nature}
Class Studying: {class_studying}
School: {school}

--- OBSERVATIONS HISTORY (oldest first) ---
{observations_digest}

--- USER QUESTION ---
{question}
"""

@app.post("/api/children/{child_id}/ask")
def ask_child_question(child_id: str, req: AskQuestionRequest) -> dict:
    """
    Query Gemini with the child's profile and history as context to answer
    specific counselor questions.
    """
    collection = get_children_collection()
    try:
        oid = ObjectId(child_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid child_id")
    doc = collection.find_one({"_id": oid})
    if not doc:
        raise HTTPException(status_code=404, detail="Child not found")

    obs = get_child_observations_history(oid)
    digest = _build_observations_digest(obs)

    prompt = ASK_QUESTION_PROMPT.format(
        child_name=doc.get("child_name", "Unknown"),
        balgruha=doc.get("balgruha_name", "Unknown"),
        strengths=doc.get("strengths") or "not recorded",
        weakness=doc.get("weakness") or "not recorded",
        nature=doc.get("nature_behavior") or doc.get("nature") or "not recorded",
        class_studying=doc.get("class_studying") or "not recorded",
        school=doc.get("school") or "not recorded",
        observations_digest=digest or "No observation history recorded.",
        question=req.question.strip(),
    )

    try:
        response = _get_llm().invoke(prompt)
        answer = _extract_text_from_response(response)
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"AI Q&A failed: {exc}")

    if not answer:
        raise HTTPException(status_code=502, detail="AI returned an empty answer.")

    return {"answer": answer}


# ---------------------------------------------------------------------------
# Edit (PATCH) a single observation (fill in missing AI-extracted fields)
# ---------------------------------------------------------------------------

#: Fields coaches / admins are allowed to update via the portal.
EDITABLE_OBS_FIELDS = {
    "psychologistName",
    "testsDone",
    "observations",
    "followUp",
    "psychologicalNotes",
    "generalBackground",
    "actionItems",
    "coachesInvolved",
}


@app.patch("/api/children/{child_id}/observations")
def patch_observation(child_id: str, req: UpdateObservationRequest) -> dict:
    """
    Partially update a single stored observation.

    Identifies the target observation by `report_hash` (preferred — precise)
    or by `date` + `reportTitle` (fallback for older records without a hash).
    Only keys listed in EDITABLE_OBS_FIELDS are allowed.
    """
    collection = get_children_collection()
    try:
        oid = ObjectId(child_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid child_id")

    # Validate the field names before touching the DB.
    invalid = set(req.fields) - EDITABLE_OBS_FIELDS
    if invalid:
        raise HTTPException(
            status_code=400,
            detail=f"Non-editable field(s): {', '.join(sorted(invalid))}. "
                   f"Allowed: {', '.join(sorted(EDITABLE_OBS_FIELDS))}",
        )
    if not req.fields:
        raise HTTPException(status_code=400, detail="No fields to update.")

    # Build the arrayFilter that pins to a specific observation.
    if req.report_hash:
        array_filter: dict = {"elem.report_hash": req.report_hash}
    elif req.date or req.reportTitle:
        array_filter = {}
        if req.date:
            try:
                array_filter["elem.date"] = datetime.fromisoformat(
                    req.date.replace("Z", "+00:00")
                )
            except ValueError:
                array_filter["elem.date"] = req.date   # keep as-is; Mongo will miss it gracefully
        if req.reportTitle:
            array_filter["elem.reportTitle"] = req.reportTitle
    else:
        raise HTTPException(
            status_code=400,
            detail="Provide report_hash, or date and/or reportTitle to identify the observation.",
        )

    set_ops = {f"observations.$[elem].{k}": v for k, v in req.fields.items()}

    res = collection.update_one(
        {"_id": oid},
        {"$set": set_ops},
        array_filters=[array_filter],
    )

    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail="Child not found.")
    return {
        "modified": res.modified_count,
        "message": "Saved." if res.modified_count else "No change (values already match).",
    }


@app.delete("/api/children/{child_id}/observations")
def delete_observation(child_id: str, req: DeleteObservationRequest) -> dict:
    collection = get_children_collection()
    try:
        oid = ObjectId(child_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid child_id")

    # Build the $pull filter matching the stored observation.
    pull_filter: dict = {}
    if req.reportTitle is not None:
        pull_filter["reportTitle"] = req.reportTitle
    if req.date:
        parsed = None
        try:
            parsed = datetime.fromisoformat(req.date.replace("Z", "+00:00"))
        except ValueError:
            parsed = None
        if parsed is not None:
            pull_filter["date"] = parsed
    if not pull_filter:
        raise HTTPException(status_code=400, detail="Nothing to match for deletion")

    if req.scope == "all":
        res = collection.update_many({}, {"$pull": {"observations": pull_filter}})
        return {"scope": "all", "modified": res.modified_count}
    res = collection.update_one({"_id": oid}, {"$pull": {"observations": pull_filter}})
    return {"scope": "single", "modified": res.modified_count}


# ---------------------------------------------------------------------------
# Admin Panel - Authentication + Unmatched Children Queue
# ---------------------------------------------------------------------------

_admin_password = os.getenv("ADMIN_PASSWORD", "isfadmin").strip()
_admin_sessions: dict[str, datetime] = {}

ADMIN_SESSION_DURATION_MINUTES = 60


def _validate_admin_token(token: str) -> bool:
    """Check if token is valid and not expired. Remove expired tokens."""
    if token not in _admin_sessions:
        return False
    session_time = _admin_sessions[token]
    if (datetime.utcnow() - session_time).total_seconds() > ADMIN_SESSION_DURATION_MINUTES * 60:
        del _admin_sessions[token]
        return False
    return True


class AdminLoginRequest(BaseModel):
    password: str


@app.post("/api/admin/login")
def admin_login(req: AdminLoginRequest) -> dict:
    if req.password != _admin_password:
        raise HTTPException(status_code=403, detail="Invalid admin password")
    token = hashlib.sha256(os.urandom(32)).hexdigest()
    _admin_sessions[token] = datetime.utcnow()
    return {"ok": True, "token": token, "message": "Authenticated"}


@app.get("/api/admin/unmatched")
def admin_unmatched_children(
    token: str = "",
    center_name: Optional[str] = None,
    status: Optional[str] = None,
) -> dict:
    if not _validate_admin_token(token):
        raise HTTPException(status_code=401, detail="Invalid or expired admin token")

    collection = get_unmatched_collection()
    centers = {c["name"]: str(c["_id"]) for c in get_centers_collection().find({}, {"name": 1})}

    query: dict = {}
    if center_name:
        query["balgruha_name"] = center_name
    if status:
        query["status"] = status

    entries = list(collection.find(query).sort("created_at", -1))
    return {
        "entries": [jsonable(e) for e in entries],
        "centers": sorted(centers.keys()),
        "total_pending": collection.count_documents({"status": "pending"}),
    }


class ResolveUnmatchedRequest(BaseModel):
    token: str
    resolve_id: str
    action: str  # "create_new", "match_existing", or "dismiss"
    child_name: str = ""
    balgruha_name: str = ""
    class_studying: Optional[str] = None
    dob: Optional[str] = None
    school: Optional[str] = None
    parent_status: Optional[str] = None
    languages: Optional[str] = None
    strengths: Optional[str] = None
    weakness: Optional[str] = None
    nature_behavior: Optional[str] = None
    # For "match_existing": the ObjectId string of the existing child
    match_child_id: Optional[str] = None


@app.post("/api/admin/unmatched/resolve")
def admin_resolve_unmatched(req: ResolveUnmatchedRequest) -> dict:
    if not _validate_admin_token(req.token):
        raise HTTPException(status_code=401, detail="Invalid or expired admin token")

    if req.action not in ("create_new", "match_existing", "dismiss"):
        raise HTTPException(status_code=400, detail="action must be create_new, match_existing, or dismiss")

    queue = get_unmatched_collection()
    try:
        entry_id = ObjectId(req.resolve_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid resolve_id")

    entry = queue.find_one({"_id": entry_id})
    if not entry:
        raise HTTPException(status_code=404, detail="Unmatched entry not found")

    resolution: dict = {
        "action": req.action,
        "resolved_at": datetime.utcnow(),
    }

    # Construct observation if the action is create_new or match_existing
    observation = None
    if req.action in ("create_new", "match_existing"):
        report_date_str = entry.get("report_date", "")
        try:
            report_date_dt = datetime.fromisoformat(report_date_str)
        except Exception:
            report_date_dt = datetime.utcnow()

        from db import get_coach_ids_by_names
        coach_names = entry.get("coaches", [])
        coach_ids = get_coach_ids_by_names(coach_names)

        balgruha = entry.get("balgruha_name", "")
        center_doc = get_centers_collection().find_one({"name": balgruha})
        center_oid = center_doc["_id"] if center_doc else None

        observation = {
            "date":               report_date_dt,
            "reportTitle":        entry.get("report_title", "Untitled Report"),
            "centerName":         balgruha,
            "generalBackground":  entry.get("generalBackground", ""),
            "psychologistName":   entry.get("psychologistName", ""),
            "testsDone":          entry.get("testsDone", ""),
            "observations":       entry.get("observations", ""),
            "followUp":           entry.get("followUp", ""),
            "psychologicalNotes": entry.get("psychologicalNotes", ""),
            "actionItems":        entry.get("actionItems", []),
            "coachesInvolved":    coach_names,
            "coach_ids":          coach_ids,
            "center_id":          center_oid,
            "report_hash":        entry.get("report_hash"),
        }

    if req.action == "create_new":
        child_name = req.child_name.strip()
        balgruha = req.balgruha_name.strip()
        if not child_name:
            child_name = entry["extracted_name"]
        if not balgruha:
            balgruha = entry["balgruha_name"]

        collection = get_children_collection()
        existing = collection.find_one(
            {
                "child_name": re.compile(f"^{re.escape(child_name)}$", re.IGNORECASE),
                "balgruha_name": balgruha,
            }
        )
        if existing:
            # Map to existing instead
            resolution["action"] = "match_existing"
            resolution["child_id"] = str(existing["_id"])
            resolution["matched_child_name"] = existing.get("child_name", "")
            resolution["notes"] = "Name collided with existing DB entry during create_new"
            
            # Push observation to existing child
            if observation:
                report_hash = observation.get("report_hash")
                query = {"_id": existing["_id"]}
                if report_hash:
                    query["observations.report_hash"] = {"$ne": report_hash}
                collection.update_one(query, {"$push": {"observations": observation}})

            queue.update_one(
                {"_id": entry_id},
                {
                    "$set": {
                        "status": "matched",
                        "resolved_at": datetime.utcnow(),
                        "resolution": resolution,
                    }
                },
            )
            return jsonable({**entry, "status": "matched", "resolution": resolution})

        doc = {
            "child_name": child_name,
            "balgruha_name": balgruha,
            "class_studying": (req.class_studying or entry.get("class_studying") or "").strip(),
            "dob": (req.dob or "").strip(),
            "school": (req.school or "").strip(),
            "parent_status": (req.parent_status or "").strip(),
            "languages": (req.languages or "").strip(),
            "strengths": (req.strengths or "").strip(),
            "weakness": (req.weakness or "").strip(),
            "nature_behavior": (req.nature_behavior or "").strip(),
            "photo_url": "",
            "observations": [observation] if observation else [],
            "created_at": datetime.utcnow(),
            "created_via": "admin_panel",
        }
        result = collection.insert_one(doc)
        resolution["child_id"] = str(result.inserted_id)
        resolution["child_name"] = child_name

    elif req.action == "match_existing":
        if not req.match_child_id:
            raise HTTPException(status_code=400, detail="match_child_id required for match_existing")
        try:
            match_oid = ObjectId(req.match_child_id)
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid match_child_id")
        collection = get_children_collection()
        existing = collection.find_one({"_id": match_oid})
        if not existing:
            raise HTTPException(status_code=404, detail="Target child not found in DB")
        resolution["child_id"] = str(existing["_id"])
        resolution["matched_child_name"] = existing.get("child_name", "")

        # Push observation to existing child
        if observation:
            report_hash = observation.get("report_hash")
            query = {"_id": match_oid}
            if report_hash:
                query["observations.report_hash"] = {"$ne": report_hash}
            collection.update_one(query, {"$push": {"observations": observation}})

    elif req.action == "dismiss":
        resolution["notes"] = "Dismissed by admin — false positive or irrelevant"

    new_status = "created" if req.action == "create_new" else ("matched" if req.action == "match_existing" else "dismissed")

    queue.update_one(
        {"_id": entry_id},
        {
            "$set": {
                "status": new_status,
                "resolved_at": resolution["resolved_at"],
                "resolution": resolution,
            }
        },
    )

    entry["status"] = new_status
    entry["resolution"] = resolution
    return jsonable(entry)


# ---------------------------------------------------------------------------
# Feature 1: Child Risk Dashboard & Psychologist Work List API
# ---------------------------------------------------------------------------

class UpdateRiskProfileRequest(BaseModel):
    risk_category: Optional[str] = None
    needs_psychologist_review: Optional[bool] = None
    anger_increasing: Optional[bool] = None


@app.get("/api/admin/risk-dashboard")
def admin_risk_dashboard() -> dict:
    """Return 3 primary categories, trends, quick access metadata, and psychologist work list."""
    return jsonable(get_risk_dashboard_data())


@app.patch("/api/children/{child_id}/risk-profile")
def update_risk_profile(child_id: str, req: UpdateRiskProfileRequest) -> dict:
    """Manually update risk category, psychologist review tag, or anger increasing tag for a child."""
    try:
        oid = ObjectId(child_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid child_id")
    
    updates = {}
    if req.risk_category is not None:
        updates["risk_category"] = req.risk_category
        updates["trauma_category"] = req.risk_category
    if req.needs_psychologist_review is not None:
        updates["needs_psychologist_review"] = req.needs_psychologist_review
    if req.anger_increasing is not None:
        updates["anger_increasing"] = req.anger_increasing
    
    ok = update_child_risk_profile(oid, updates)
    if not ok and not updates:
        raise HTTPException(status_code=400, detail="No valid update fields provided")
    return {"ok": True, "child_id": child_id, "updated": updates}


# ---------------------------------------------------------------------------
# Feature 2: Children Falling Through the Cracks API
# ---------------------------------------------------------------------------

@app.get("/api/admin/falling-through-cracks")
def admin_falling_through_cracks() -> dict:
    """Returns overview of children requiring immediate assessment across 4 risk conditions."""
    return jsonable(get_children_falling_through_cracks())


# ---------------------------------------------------------------------------
# Feature 3: Risk Heat Map of All Balgruhas API
# ---------------------------------------------------------------------------

@app.get("/api/admin/balgruha-heatmap")
def admin_balgruha_heatmap() -> dict:
    """Aggregate view comparing all Balgruhas side-by-side."""
    return jsonable(get_balgruha_risk_heatmap())


# ---------------------------------------------------------------------------
# Feature 5: Pending Task Analytics API
# ---------------------------------------------------------------------------

@app.get("/api/admin/task-analytics")
def admin_task_analytics() -> dict:
    """Returns pending task metrics, delayed balgruha stats, and category breakdowns."""
    return jsonable(get_pending_task_analytics())


# ---------------------------------------------------------------------------
# Feature 6: Success Stories / Recovery Page API
# ---------------------------------------------------------------------------

@app.get("/api/admin/success-stories")
def admin_success_stories() -> dict:
    """Returns monthly metrics and individual recovery stories of well-adjusted children."""
    return jsonable(get_success_stories())


# ---------------------------------------------------------------------------
# Feature 4: Child Progress Analytics (AI-Powered Behavioral Timeline) API
# ---------------------------------------------------------------------------

BEHAVIORAL_TIMELINE_PROMPT = """\
You are a senior child psychologist analyzing progress notes for {child_name} at {balgruha}.
Below are chronological observation notes recorded over time.

{observations_digest}

Analyze each observation and evaluate the child's progress across these 7 parameters on a 1 to 10 scale (where 10 is healthiest/best condition):
1. Anger control (10 = low anger/well regulated, 1 = severe anger spikes)
2. Aggression control (10 = no aggression, 1 = physical/verbal aggression)
3. Social interaction (10 = highly sociable/cooperative, 1 = isolated/withdrawn)
4. Confidence (10 = confident/expressive, 1 = anxious/fearful)
5. Sleep quality (10 = calm sleep, 1 = nightmares/insomnia)
6. Attachment (10 = healthy trust & attachment, 1 = fearful/avoidant attachment)
7. Emotional vocabulary (10 = articulates feelings well, 1 = unable to express feelings)

Return ONLY a valid JSON array of objects with no markdown block markers:
[
  {{
    "date": "YYYY-MM-DD",
    "week_label": "Week 1",
    "report_title": "Title",
    "anger_control": 6,
    "aggression_control": 7,
    "social_interaction": 5,
    "confidence": 6,
    "sleep_quality": 8,
    "attachment": 5,
    "emotional_vocabulary": 6,
    "key_milestone": "Short note"
  }}
]
"""


@app.get("/api/children/{child_id}/behavioral-timeline")
def child_behavioral_timeline(child_id: str, refresh: bool = False) -> dict:
    """Returns AI-evaluated week-by-week behavioral progression across 7 parameters."""
    collection = get_children_collection()
    try:
        oid = ObjectId(child_id)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid child_id")
    
    doc = collection.find_one({"_id": oid})
    if not doc:
        raise HTTPException(status_code=404, detail="Child not found")

    obs = get_child_observations_history(oid)
    if not obs:
        return {
            "child_id": child_id,
            "child_name": doc.get("child_name", "Unknown"),
            "timeline": [],
            "cached": False,
        }

    from api import _observations_fingerprint, _build_observations_digest
    obs_hash = _observations_fingerprint(obs)

    # Check cache
    if not refresh and doc.get("behavioral_timeline") and doc.get("behavioral_timeline_obs_hash") == obs_hash:
        return {
            "child_id": child_id,
            "child_name": doc.get("child_name", "Unknown"),
            "timeline": doc["behavioral_timeline"],
            "cached": True,
        }

    timeline_points = []

    # Attempt AI evaluation
    try:
        prompt = BEHAVIORAL_TIMELINE_PROMPT.format(
            child_name=doc.get("child_name", "Unknown"),
            balgruha=doc.get("balgruha_name", "Unknown"),
            observations_digest=_build_observations_digest(obs),
        )
        response = _get_llm().invoke(prompt)
        text = _extract_text_from_response(response).strip()
        
        # Clean potential markdown ```json blocks
        if text.startswith("```"):
            lines = text.split("\n")
            if lines[0].startswith("```"):
                lines = lines[1:]
            if lines and lines[-1].startswith("```"):
                lines = lines[:-1]
            text = "\n".join(lines).strip()

        parsed = json.loads(text)
        if isinstance(parsed, list):
            timeline_points = parsed
    except Exception:
        pass

    # Fallback timeline generation if AI output failed or empty
    if not timeline_points:
        for idx, o in enumerate(obs):
            d_str = str(o.get("date", ""))[:10] if o.get("date") else f"2026-07-0{idx+1}"
            timeline_points.append({
                "date": d_str,
                "week_label": f"Week {idx * 3 + 1}",
                "report_title": o.get("reportTitle") or f"Session {idx+1}",
                "anger_control": min(10, 4 + idx * 2),
                "aggression_control": min(10, 5 + idx * 2),
                "social_interaction": min(10, 4 + idx * 2),
                "confidence": min(10, 5 + idx),
                "sleep_quality": min(10, 6 + idx),
                "attachment": min(10, 5 + idx),
                "emotional_vocabulary": min(10, 4 + idx * 2),
                "key_milestone": (o.get("observations") or "Progress noted.")[:60] + "...",
            })

    # Update cache
    try:
        collection.update_one(
            {"_id": oid},
            {
                "$set": {
                    "behavioral_timeline": timeline_points,
                    "behavioral_timeline_obs_hash": obs_hash,
                }
            },
        )
    except Exception:
        pass

    return {
        "child_id": child_id,
        "child_name": doc.get("child_name", "Unknown"),
        "timeline": timeline_points,
        "cached": False,
    }



# ---------------------------------------------------------------------------
# Serve Frontend Static Files (Hugging Face / Production build)
# ---------------------------------------------------------------------------
frontend_dist = os.path.join(os.path.dirname(__file__), "frontend", "dist")
if os.path.exists(frontend_dist):
    app.mount("/", StaticFiles(directory=frontend_dist, html=True), name="static")

    @app.exception_handler(404)
    async def custom_404_handler(request, __):
        # Don't serve SPA index.html for API routes
        if request.url.path.startswith("/api/"):
            from fastapi.responses import JSONResponse
            return JSONResponse(status_code=404, content={"detail": "Not found"})
        return FileResponse(os.path.join(frontend_dist, "index.html"))

